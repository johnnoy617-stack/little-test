import re
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

from app.extensions import db
from app.models import Chunk, Document


class DocumentProcessingError(Exception):
    pass


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def save_and_parse_document(app, uploaded_file: FileStorage) -> tuple[Document, list[Chunk]]:
    if not uploaded_file or not uploaded_file.filename:
        raise DocumentProcessingError("请选择要上传的 PDF 或 DOCX 文件。")

    original_name = uploaded_file.filename
    suffix = Path(original_name).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentProcessingError("仅支持 PDF 和 DOCX，DOC 请先另存为 DOCX。")

    safe_name = secure_filename(original_name) or f"document{suffix}"
    generated_name = f"{uuid.uuid4().hex}_{safe_name}"
    file_path = Path(app.config["UPLOAD_DIR"]) / generated_name
    uploaded_file.save(file_path)

    try:
        units = extract_document_units(file_path, suffix)
        if not units:
            raise DocumentProcessingError("没有读取到可用文本内容，请确认文件不是扫描件或空文档。")

        chunk_payloads = chunk_units(
            units,
            chunk_size=app.config["CHUNK_SIZE"],
            overlap=app.config["CHUNK_OVERLAP"],
        )
        if not chunk_payloads:
            raise DocumentProcessingError("没有生成可用知识片段，请检查文档内容。")

        document = Document(
            filename=generated_name,
            original_name=original_name,
            file_type=suffix.lstrip("."),
            storage_path=str(file_path),
            page_count=len(units),
            chunk_count=len(chunk_payloads),
            status="processing",
        )
        db.session.add(document)
        db.session.flush()

        chunks: list[Chunk] = []
        for payload in chunk_payloads:
            chunk = Chunk(
                document_id=document.id,
                page_number=payload["page_number"],
                position_label=payload["position_label"],
                chunk_index=payload["chunk_index"],
                content=payload["content"],
                char_count=len(payload["content"]),
            )
            db.session.add(chunk)
            chunks.append(chunk)

        db.session.flush()
        return document, chunks
    except Exception:
        db.session.rollback()
        if file_path.exists():
            file_path.unlink()
        raise


def finalize_document(document: Document, chunks: list[Chunk], embedding_model: str) -> None:
    for chunk in chunks:
        chunk.qdrant_point_id = str(chunk.id)
        chunk.embedding_model = embedding_model

    document.status = "ready"
    document.error_message = None
    document.chunk_count = len(chunks)
    db.session.commit()


def mark_document_failed(document: Document, message: str) -> None:
    document.status = "failed"
    document.error_message = message
    db.session.commit()


def delete_document_file(document: Document) -> None:
    if document.storage_path:
        file_path = Path(document.storage_path)
    else:
        file_path = Path(document.filename)

    if file_path.exists():
        file_path.unlink()


def delete_document_record(document: Document) -> None:
    db.session.delete(document)
    db.session.commit()


def extract_document_units(file_path: Path, suffix: str) -> list[dict]:
    if suffix == ".pdf":
        return extract_pdf_units(file_path)
    if suffix == ".docx":
        return extract_docx_units(file_path)
    raise DocumentProcessingError("暂不支持该文件类型。")


def extract_pdf_units(file_path: Path) -> list[dict]:
    try:
        reader = PdfReader(str(file_path))
    except Exception as exc:
        raise DocumentProcessingError("PDF 无法解析，请确认文件未损坏。") from exc

    units: list[dict] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = normalize_text(page.extract_text() or "")
        if text:
            units.append(
                {
                    "page_number": page_index,
                    "position_label": f"Page {page_index}",
                    "text": text,
                }
            )
    return units


def extract_docx_units(file_path: Path) -> list[dict]:
    try:
        doc = DocxDocument(str(file_path))
    except Exception as exc:
        raise DocumentProcessingError("DOCX 无法解析，请确认文件格式正确。") from exc

    units: list[dict] = []
    paragraph_index = 1
    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text or "")
        if not text:
            continue
        units.append(
            {
                "page_number": paragraph_index,
                "position_label": f"Paragraph {paragraph_index}",
                "text": text,
            }
        )
        paragraph_index += 1
    return units


def chunk_units(units: list[dict], chunk_size: int = 700, overlap: int = 120) -> list[dict]:
    chunks: list[dict] = []
    effective_overlap = min(overlap, max(0, chunk_size - 50))

    for unit in units:
        text = unit["text"]
        start = 0
        chunk_index = 0

        while start < len(text):
            end = min(start + chunk_size, len(text))
            content = text[start:end].strip()
            if content:
                chunks.append(
                    {
                        "page_number": unit["page_number"],
                        "position_label": unit["position_label"],
                        "chunk_index": chunk_index,
                        "content": content,
                    }
                )
                chunk_index += 1

            if end >= len(text):
                break
            start = max(end - effective_overlap, start + 1)

    return chunks


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()
